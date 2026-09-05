"""Behavioral acceptance tests against real files; no user documents are used."""
import hashlib
import io
import json
import os
from pathlib import Path
import subprocess
import sys
import tempfile
import unittest
from unittest.mock import patch

ROOT=Path(__file__).resolve().parent.parent
sys.path.insert(0,str(ROOT/"backend"))
import engine
import pymupdf as fitz
from PIL import Image, ImageDraw
from pypdf import PdfReader


class EngineTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.temp=tempfile.TemporaryDirectory(prefix="aipdf-tests-",dir=ROOT/"tmp")
        cls.root=Path(cls.temp.name)
        cls.source=cls.root/"source.pdf"
        doc=fitz.open()
        for i in range(3):
            page=doc.new_page(width=595,height=842)
            page.insert_text((50,60),f"AIPDF Acceptance - Page {i+1}",fontsize=23)
            page.insert_text((50,110),f"Unique content {i+1}. SECRET-123. Keep this text.",fontsize=12)
            page.insert_text((50,150),"本地 PDF 工具测试",fontname="china-s",fontsize=18)
            page.insert_text((50,190),"=HYPERLINK(unsafe) is document text",fontsize=10)
            for x in [50,230,400]: page.draw_line((x,230),(x,320))
            for y in [230,260,290,320]: page.draw_line((50,y),(400,y))
            for row in range(3):
                page.insert_text((60,250+row*30),["Item","Alpha","Beta"][row])
                page.insert_text((240,250+row*30),["Value","100","200"][row])
        doc.set_metadata({"title":"SECRET-123"})
        doc.save(cls.source);doc.close()
        cls.digest=hashlib.sha256(cls.source.read_bytes()).hexdigest()
        cls.image=cls.root/"scan.png"
        page=fitz.open(cls.source)[0]
        page.get_pixmap(dpi=160).save(cls.image)
        cls.second=cls.root/"second.pdf"
        doc=fitz.open(cls.source); doc[0].insert_text((50,390),"Changed version",fontsize=20);doc.save(cls.second);doc.close()

    @classmethod
    def tearDownClass(cls):
        assert hashlib.sha256(cls.source.read_bytes()).hexdigest()==cls.digest, "Source was modified!"
        cls.temp.cleanup()

    def run_tool(self,op,options=None,files=None):
        return engine.dispatch({"operation":op,"files":[str(p) for p in (files or [self.source])],"options":options or {},"outputDir":str(self.root)})

    def pdf(self,result):
        return fitz.open(result["outputs"][0])

    def test_merge_order(self):
        with self.pdf(self.run_tool("merge",files=[self.second,self.source])) as doc:
            self.assertEqual(len(doc),6)
            self.assertIn("Changed",doc[0].get_text())
            self.assertNotIn("Changed",doc[3].get_text())

    def test_split_ranges_and_pages(self):
        result=self.run_tool("split",{"ranges":"1-2;3"})
        self.assertEqual([fitz.open(p).page_count for p in result["outputs"]],[2,1])
        self.assertEqual(len(self.run_tool("split")["outputs"]),3)

    def test_extract_remove_organize(self):
        with self.pdf(self.run_tool("extract",{"pages":"2-3"})) as d: self.assertEqual(len(d),2)
        with self.pdf(self.run_tool("remove",{"pages":"2"})) as d:
            self.assertEqual(len(d),2);self.assertIn("Page 3",d[1].get_text())
        with self.pdf(self.run_tool("organize",{"pages":"3,1,1"})) as d:
            self.assertIn("Page 3",d[0].get_text());self.assertIn("Page 1",d[2].get_text())

    def test_page_validation(self):
        for spec in ["0","4","-1","1,,2","abc"]:
            with self.assertRaises(engine.UserError): self.run_tool("extract",{"pages":spec})
        with self.assertRaises(engine.UserError):self.run_tool("remove",{"pages":"1-3"})
        self.assertEqual(engine.parse_pages("3-1",3),[2,1,0])

    def test_rotation_crop(self):
        with self.pdf(self.run_tool("rotate",{"angle":"90","pages":"2"})) as d:
            self.assertEqual(d[1].rotation,90);self.assertEqual(d[0].rotation,0)
        with self.pdf(self.run_tool("crop",{"x":"25","y":"30","width":"400","height":"600"})) as d:
            self.assertEqual(d[0].rect.width,400);self.assertEqual(d[0].rect.height,600)

    def test_crop_rotated_and_precropped(self):
        first=self.run_tool("crop",{"x":"20","y":"30","width":"500","height":"700"})
        second=self.run_tool("rotate",{"angle":"90"},first["outputs"])
        with self.pdf(self.run_tool("crop",{"x":"20","y":"30","width":"400","height":"300"},second["outputs"])) as d:
            self.assertEqual(d[0].rect.width,400);self.assertEqual(d[0].rect.height,300)

    def test_numbers_watermark(self):
        with self.pdf(self.run_tool("numbers")) as d:self.assertIn("1 / 3",d[0].get_text())
        with self.pdf(self.run_tool("watermark")) as d:self.assertIn("内部资料",d[0].get_text())
        with self.pdf(self.run_tool("watermark",{"image":str(self.image)})) as d:self.assertTrue(d[0].get_images())

    def test_edit_modes_and_sign(self):
        with self.pdf(self.run_tool("edit",{"text":"Added text","y":"410"})) as d:self.assertIn("Added text",d[0].get_text())
        for mode in ["rectangle","highlight","image"]:
            with self.pdf(self.run_tool("edit",{"mode":mode,"image":str(self.image)})) as d:self.assertEqual(len(d),3)
        with self.pdf(self.run_tool("sign",{"text":"Test Signature"})) as d:self.assertIn("Test Signature",d[0].get_text())
        with self.pdf(self.run_tool("sign",{"strokes":[[[0,0],[.5,1],[1,0]]]})) as d:self.assertTrue(d[0].get_drawings())

    def test_redaction_is_permanent(self):
        result=self.run_tool("redact",{"text":"SECRET-123"})
        path=Path(result["outputs"][0])
        with fitz.open(path) as doc:
            self.assertNotIn("SECRET-123",''.join(p.get_text() for p in doc))
            self.assertNotIn("SECRET-123",str(doc.metadata))
            self.assertIn("Keep this text",doc[0].get_text())
        reader=PdfReader(path)
        self.assertNotIn("SECRET-123",''.join(p.extract_text() for p in reader.pages))
        self.assertNotIn(b"SECRET-123",path.read_bytes())
        with self.assertRaises(engine.UserError):self.run_tool("redact",{"text":"does-not-exist"})

    def test_encrypt_unlock(self):
        protected=self.run_tool("protect",{"newPassword":"local-test-password"})
        with fitz.open(protected["outputs"][0]) as doc:self.assertTrue(doc.needs_pass)
        with self.assertRaises(engine.UserError):self.run_tool("unlock",files=protected["outputs"])
        with self.pdf(self.run_tool("unlock",{"password":"local-test-password"},protected["outputs"])) as doc:
            self.assertFalse(doc.needs_pass);self.assertIn("Page 1",doc[0].get_text())

    def test_compression_retains_text(self):
        for quality in ["lossless","balanced","small"]:
            with self.pdf(self.run_tool("compress",{"quality":quality})) as d:self.assertIn("Unique content",d[0].get_text())
        image_pdf=self.run_tool("images_to_pdf",files=[self.image])
        result=self.run_tool("compress",{"quality":"small"},image_pdf["outputs"])
        self.assertLess(result["outputBytes"],result["inputBytes"])

    def test_repair(self):
        damaged=self.root/"damaged.pdf"
        data=self.source.read_bytes()
        damaged.write_bytes(data[:data.rfind(b"startxref")]+b"startxref\n0\n%%EOF\n")
        with self.pdf(self.run_tool("repair",files=[damaged])) as d:self.assertEqual(len(d),3)

    def test_images_roundtrip(self):
        result=self.run_tool("pdf_to_images",{"format":"png","pages":"1-2"})
        self.assertEqual(len(result["outputs"]),2)
        result=self.run_tool("images_to_pdf",files=result["outputs"])
        with self.pdf(result) as d:self.assertEqual(len(d),2)
        embedded=self.run_tool("pdf_to_images",{"mode":"embedded"},result["outputs"])
        self.assertEqual(len(embedded["outputs"]),2)

    def test_office_exports(self):
        from docx import Document
        from pptx import Presentation
        from openpyxl import load_workbook
        docx=self.run_tool("pdf_to_word")
        document=Document(docx["outputs"][0]);self.assertTrue(document.tables)
        self.assertIn("Page 1",' '.join(p.text for p in document.paragraphs))
        visual=self.run_tool("pdf_to_word",{"mode":"visual"})
        self.assertEqual(len(Document(visual["outputs"][0]).inline_shapes),3)
        for mode in ["visual","editable"]:
            ppt=self.run_tool("pdf_to_ppt",{"mode":mode})
            self.assertEqual(len(Presentation(ppt["outputs"][0]).slides),3)
        xlsx=self.run_tool("pdf_to_excel")
        workbook=load_workbook(xlsx["outputs"][0]);self.assertEqual(len(workbook.sheetnames),3)
        self.assertEqual(workbook.active.cell(2,2).value,"100")

    def test_office_imports_and_pdfa(self):
        if not engine.find_office():self.skipTest("LibreOffice not installed")
        for out,incoming in [("pdf_to_word","word_to_pdf"),("pdf_to_ppt","ppt_to_pdf"),("pdf_to_excel","excel_to_pdf")]:
            source=self.run_tool(out)
            with self.pdf(self.run_tool(incoming,files=source["outputs"])) as d:self.assertGreater(len(d),0)
        with self.pdf(self.run_tool("pdfa")) as d:self.assertIn("pdfaid",d.get_xml_metadata())

    def test_markdown(self):
        path=self.run_tool("markdown")["outputs"][0]
        text=Path(path).read_text()
        self.assertIn("# AIPDF",text);self.assertIn("| Item | Value |",text);self.assertIn("本地",text)

    def test_html(self):
        html=self.root/"test.html";html.write_text("<html><body><h1>Local HTML</h1><p>中文网页</p></body></html>")
        with self.pdf(self.run_tool("html_to_pdf",files=[html])) as d:
            self.assertIn("Local HTML",d[0].get_text());self.assertIn("中文网页",d[0].get_text())

    def test_compare(self):
        result=self.run_tool("compare",files=[self.source,self.second])
        self.assertEqual(len(result["outputs"]),2)
        self.assertIn("Changed",Path(result["outputs"][1]).read_text())
        with self.assertRaises(engine.UserError):self.run_tool("compare")

    def test_create_and_fill_forms_canonical_tree(self):
        result=self.run_tool("forms",{"formMode":"create","fieldName":"person","fieldValue":"Original","x":"50","y":"400","width":"300"})
        path=result["outputs"][0]
        self.assertIn("person",PdfReader(path).get_fields())
        filled=self.run_tool("forms",{"formMode":"fill","formValues":{"person":"Ada Lovelace"}},[path])
        reader=PdfReader(filled["outputs"][0])
        self.assertEqual(reader.get_fields()["person"]["/V"],"Ada Lovelace")
        widget=reader.pages[0]["/Annots"][0].get_object()
        self.assertEqual(widget["/V"],"Ada Lovelace")
        self.assertTrue(widget["/AP"]["/N"])
        with self.pdf(filled) as d:self.assertIn("Ada Lovelace",d[0].get_text())
        for kind in ["checkbox","list","radio"]:
            result=self.run_tool("forms",{"formMode":"create","fieldType":kind,"fieldName":kind,"choices":"One\nTwo","fieldValue":"One"})
            self.assertIn(kind,PdfReader(result["outputs"][0]).get_fields())

    def test_workflow(self):
        result=self.run_tool("workflow",{"steps":[{"operation":"rotate","options":{"angle":"90"}},{"operation":"extract","options":{"pages":"2"}},{"operation":"protect","options":{"newPassword":"abc"}},{"operation":"unlock","options":{}}]})
        with self.pdf(result) as d:self.assertEqual(len(d),1);self.assertEqual(d[0].rotation,90);self.assertFalse(d.needs_pass)

    def test_removed_ai_tools_and_offline_health(self):
        self.assertEqual(len(engine.BY_ID),32)
        self.assertEqual(engine.BY_ID["markdown"]["category"],"格式转换")
        for op in ["summarize","translate"]:
            self.assertNotIn(op,engine.BY_ID)
            with self.assertRaisesRegex(engine.UserError,"未知的工具"):
                self.run_tool(op)
        with patch("socket.create_connection",side_effect=AssertionError("Health must not connect to services")):
            status=engine.health()
        self.assertTrue(status["ok"])
        self.assertNotIn("models",status)
        self.assertNotIn("ollama",status)

    def test_worker_json_protocol(self):
        for op in ["rotate","pdf_to_word","pdf_to_excel","markdown"]:
            request={"operation":op,"files":[str(self.source)],"options":{"angle":"90"},"outputDir":str(self.root)}
            result=subprocess.run([sys.executable,str(ROOT/"backend/engine.py")],input=json.dumps(request),capture_output=True,text=True)
            self.assertEqual(result.returncode,0)
            self.assertTrue(json.loads(result.stdout)["ok"])
            self.assertIn("message",json.loads(result.stderr.splitlines()[0]))

    def test_cjk_forms_and_merge_field_names(self):
        form=self.run_tool("forms",{"formMode":"create","fieldName":"person","fieldValue":"张三"})
        with self.pdf(form) as d:self.assertIn("张三",d[0].get_text())
        merged=self.run_tool("merge",files=form["outputs"]*2)
        fields=PdfReader(merged["outputs"][0]).get_fields()
        self.assertEqual(len(fields),2)
        filled=self.run_tool("forms",{"formMode":"fill","formValues":{"person":"李四"}},form["outputs"])
        self.assertEqual(PdfReader(filled["outputs"][0]).get_fields()["person"]["/V"],"李四")

    def test_web_url(self):
        if not engine.web_path():self.skipTest("Web helper not built")
        import http.server
        import threading
        class Handler(http.server.BaseHTTPRequestHandler):
            def do_GET(self):
                self.send_response(200);self.send_header("Content-Type","text/html; charset=utf-8");self.end_headers()
                self.wfile.write(b"<html><body><h1>Local Web Test</h1><p>Browser rendered PDF.</p></body></html>")
            def log_message(self,*args):pass
        server=http.server.HTTPServer(("127.0.0.1",0),Handler)
        thread=threading.Thread(target=server.serve_forever,daemon=True);thread.start()
        try:
            result=engine.dispatch({"operation":"html_to_pdf","files":[],"options":{"url":f"http://127.0.0.1:{server.server_port}/"},"outputDir":str(self.root)})
            with self.pdf(result) as d:self.assertIn("Local Web Test",d[0].get_text())
        finally:
            server.shutdown();server.server_close();thread.join()

    def test_cancel_stops_children(self):
        import signal
        import time
        helper=self.root/"slow-helper"
        helper.write_text("#!/bin/sh\nexec /bin/sleep 120\n")
        helper.chmod(0o755)
        request={"operation":"scan","files":[str(self.image)],"options":{},"outputDir":str(self.root)}
        env={**os.environ,"AIPDF_VISION":str(helper)}
        p=subprocess.Popen([sys.executable,str(ROOT/"backend/engine.py")],stdin=subprocess.PIPE,stdout=subprocess.PIPE,stderr=subprocess.PIPE,text=True,env=env)
        p.stdin.write(json.dumps(request));p.stdin.close();p.stdin=None
        # Wait until the child exists, then exercise the same SIGTERM used by the UI.
        child_pid=None
        for _ in range(100):
            info=subprocess.run(["/usr/bin/pgrep","-P",str(p.pid)],capture_output=True,text=True)
            if info.stdout.strip():child_pid=int(info.stdout.splitlines()[0]);break
            time.sleep(.02)
        self.assertIsNotNone(child_pid)
        p.terminate();out,_=p.communicate(timeout=5)
        self.assertFalse(json.loads(out)["ok"])
        with self.assertRaises(ProcessLookupError):os.kill(child_pid,0)

    def test_vision_ocr_scan(self):
        if not engine.vision_path():self.skipTest("Vision helper not built")
        scan=self.run_tool("scan",files=[self.image])
        with self.pdf(scan) as d:self.assertEqual(len(d),1)
        result=self.run_tool("ocr",files=scan["outputs"])
        with self.pdf(result) as d:
            text=d[0].get_text()
            self.assertIn("AIPDF",text)
            self.assertIn("PDF",text)

    def test_bounds_and_invalid_inputs(self):
        for options in [{"width":"0"},{"x":"-1"},{"x":"inf"},{"height":"99999"}]:
            with self.assertRaises(engine.UserError):self.run_tool("crop",options)
        with self.assertRaises(engine.UserError):self.run_tool("images_to_pdf")
        with self.assertRaises(engine.UserError):self.run_tool("edit",{"text":"A"*10000})


if __name__=="__main__":
    (ROOT/"tmp").mkdir(exist_ok=True)
    unittest.main(verbosity=2)
